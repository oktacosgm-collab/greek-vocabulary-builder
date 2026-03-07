"""modules/word_export.py — Word document export"""
import io
from .config import ARTICLE


def build_word_doc(words_to_export: list, all_words: dict, export_opts: dict = None) -> io.BytesIO:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if export_opts is None:
        export_opts = {"pos": True, "translation": True, "definition": False, "example": True}

    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    title = doc.add_heading("🇬🇷 Greek Vocabulary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5c)

    sub = doc.add_paragraph(f"{len(words_to_export)} words exported")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    sub.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    for i, word in enumerate(words_to_export, 1):
        d = all_words.get(word, {})

        # Word heading — always shown
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {word}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5c)

        # POS + difficulty + gender
        if export_opts.get("pos"):
            decl = d.get("declension", {}) or {}
            gender = decl.get("gender", "") if isinstance(decl, dict) else ""
            article = ARTICLE.get(gender, "")
            gender_str = f"  ·  {article} ({gender})" if gender else ""
            meta = doc.add_paragraph()
            r2 = meta.add_run(f"{d.get('part_of_speech','—')}  ·  {d.get('difficulty','?')}{gender_str}")
            r2.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
            r2.font.size = Pt(10)

        # Translation
        if export_opts.get("translation"):
            t = doc.add_paragraph()
            t.add_run("Translation: ").bold = True
            t.add_run(d.get('translation', '—'))

        # Definition
        if export_opts.get("definition"):
            df = doc.add_paragraph()
            df.add_run("Definition: ").bold = True
            df.add_run(d.get('definition', '—'))

        # Example
        if export_opts.get("example"):
            ex_gr = d.get('example_greek', '')
            ex_en = d.get('example_english', '')
            if ex_gr:
                ex = doc.add_paragraph()
                ex.add_run("Παράδειγμα: ").bold = True
                r = ex.add_run(ex_gr)
                r.italic = True
                if ex_en:
                    en = doc.add_paragraph()
                    en_r = en.add_run(ex_en)
                    en_r.italic = True
                    en_r.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
                    en_r.font.size = Pt(10)

        if i < len(words_to_export):
            doc.add_paragraph("─" * 60).runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
